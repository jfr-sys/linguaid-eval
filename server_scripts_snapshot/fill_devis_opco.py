#!/usr/bin/env python3
# -*- coding: utf-8 -*-
# fill_devis_opco.py — genere un devis client conforme OPCO depuis
# views/DEVIS_OPCO_TEMPLATE.docx (en-tete/pied de page officiels Linguaid).
# Usage: python3 fill_devis_opco.py '<json>'
# JSON in:  devisNumber, dateEmission, dateValidite, clientBlock (multiline),
#           clientName, intitule, stagiaire, modalite, periode, duree, formateur,
#           refsOpco (optional), designation, qte, pu, montant,
#           travauxLabel/travauxMontant (optional), total, outDir, id
# JSON out: {success, docxPath, pdfPath} or {success:false, error}
import sys, os, json, subprocess

TEMPLATE = '/var/www/vhosts/linguaid.net/eval.linguaid.net/app/views/DEVIS_OPCO_TEMPLATE.docx'

def main():
    try:
        args = json.loads(sys.argv[1])
    except Exception as e:
        print(json.dumps({'success': False, 'error': 'bad json: %s' % e})); return

    try:
        from docx import Document
        doc = Document(TEMPLATE)

        refs = (args.get('refsOpco') or '').strip()
        trav_label = (args.get('travauxLabel') or '').strip()
        trav_montant = (args.get('travauxMontant') or '').strip()
        has_trav = bool(trav_montant)

        mapping = {
            '{{DEVIS_NUM}}': args.get('devisNumber', ''),
            '{{DATE_EMISSION}}': args.get('dateEmission', ''),
            '{{DATE_VALIDITE}}': args.get('dateValidite', ''),
            '{{CLIENT_NAME}}': args.get('clientName', ''),
            '{{INTITULE}}': args.get('intitule', ''),
            '{{STAGIAIRE}}': args.get('stagiaire', ''),
            '{{MODALITE}}': args.get('modalite', ''),
            '{{PERIODE}}': args.get('periode', ''),
            '{{DUREE}}': args.get('duree', ''),
            '{{FORMATEUR}}': args.get('formateur', ''),
            '{{REFS_OPCO}}': refs,
            '{{DESIGNATION}}': args.get('designation', ''),
            '{{QTE}}': args.get('qte', ''),
            '{{PU}}': args.get('pu', ''),
            '{{MONTANT}}': args.get('montant', ''),
            '{{DESIGNATION2}}': trav_label,
            '{{QTE2}}': args.get('travauxQte', ''),
            '{{PU2}}': args.get('travauxPu', ''),
            '{{MONTANT2}}': trav_montant,
            '{{TOTAL}}': args.get('total', ''),
        }

        def iter_paras(container):
            for p in container.paragraphs:
                yield p
            for t in getattr(container, 'tables', []):
                for row in t.rows:
                    for cell in row.cells:
                        for p in iter_paras(cell):
                            yield p

        # 1. Expand CLIENT_BLOCK: token paragraph -> one paragraph per line,
        #    first line keeps the (bold) formatting of the token paragraph.
        import copy as _copy
        client_lines = [l.strip() for l in (args.get('clientBlock') or '').split('\n') if l.strip()]
        if not client_lines:
            print(json.dumps({'success': False, 'error': 'clientBlock vide'})); return
        target = None
        for p in iter_paras(doc):
            if '{{CLIENT_BLOCK}}' in p.text:
                target = p; break
        if target is None:
            print(json.dumps({'success': False, 'error': 'CLIENT_BLOCK token absent du template'})); return
        # first line in the token paragraph (bold company name)
        for r in target.runs:
            r.text = r.text.replace('{{CLIENT_BLOCK}}', client_lines[0])
        # subsequent lines: cloned paragraphs, non-bold
        anchor_p = target._p
        for line in client_lines[1:]:
            newp = _copy.deepcopy(anchor_p)
            anchor_p.addnext(newp)
            anchor_p = newp
            # set text on clone
            from docx.text.paragraph import Paragraph as _P
            pp = _P(newp, target._parent)
            first = True
            for r in pp.runs:
                if first:
                    r.text = line
                    r.font.bold = False
                    first = False
                else:
                    r.text = ''

        # 2. Optional rows/lines removal
        #    refsOpco empty -> remove its whole line
        if not refs:
            for p in list(iter_paras(doc)):
                if '{{REFS_OPCO}}' in p.text:
                    p._p.getparent().remove(p._p)
        #    travaux empty -> remove second item row of the price table
        price_tbl = None
        for t in doc.tables:
            if any('{{DESIGNATION2}}' in row.cells[0].text for row in t.rows):
                price_tbl = t; break
        if not has_trav and price_tbl is not None:
            for row in list(price_tbl.rows):
                if '{{DESIGNATION2}}' in row.cells[0].text:
                    row._tr.getparent().remove(row._tr)

        # 3. Run-level token replacement everywhere in body
        for p in iter_paras(doc):
            for r in p.runs:
                if '{{' in r.text:
                    for k, v in mapping.items():
                        if k in r.text:
                            r.text = r.text.replace(k, v)

        # 4. Save + convert
        out_dir = args.get('outDir') or '/tmp'
        if not os.path.isdir(out_dir):
            os.makedirs(out_dir, exist_ok=True)
        base = 'devis_opco_%s' % args.get('id', 'x')
        docx_path = os.path.join(out_dir, base + '.docx')
        pdf_path = os.path.join(out_dir, base + '.pdf')
        doc.save(docx_path)

        env = dict(os.environ)
        env['HOME'] = '/tmp'
        r = subprocess.run(['soffice', '--headless', '--convert-to', 'pdf',
                            '--outdir', out_dir, docx_path],
                           capture_output=True, timeout=120, env=env)
        if not os.path.exists(pdf_path):
            print(json.dumps({'success': False,
                              'error': 'PDF conversion failed: %s' % r.stderr.decode('utf-8', 'ignore')[:300]}))
            return

        # 5. Leftover-token safety check on the PDF text layer is overkill;
        #    check the docx XML instead.
        import zipfile, re
        xml = zipfile.ZipFile(docx_path).read('word/document.xml').decode('utf-8')
        left = re.findall(r'\{\{[A-Z0-9_]+\}\}', xml)
        if left:
            print(json.dumps({'success': False, 'error': 'tokens non remplaces: %s' % sorted(set(left))}))
            return

        print(json.dumps({'success': True, 'docxPath': docx_path, 'pdfPath': pdf_path}))
    except Exception as e:
        import traceback
        print(json.dumps({'success': False, 'error': '%s: %s' % (type(e).__name__, e),
                          'trace': traceback.format_exc()[-800:]}))

if __name__ == '__main__':
    main()
